"""Convert Markdown documentation files to Word (.docx) in documentation/."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "documentation"

DOCS = [
    (ROOT / "README.md", OUT_DIR / "Live_Seller_Churn_Analysis_Documentation.docx"),
    (
        OUT_DIR / "churn_dashboard_guide.md",
        OUT_DIR / "Live_Seller_Churn_Dashboard_Guide.docx",
    ),
]


def set_cell_shading(cell, fill_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_paragraph(doc, text, style=None, bold=False, monospace=False):
    p = doc.add_paragraph(style=style)
    if not text:
        return p
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            cleaned = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", part)
            run = p.add_run(cleaned)
            if bold:
                run.bold = True
            if monospace:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
    return p


def is_table_row(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1]


def parse_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator_row(cells):
    return all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in cells)


def add_table(doc, rows):
    if not rows:
        return
    headers = rows[0]
    data_rows = rows[1:] if len(rows) > 1 else []
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = re.sub(r"\*\*([^*]+)\*\*", r"\1", h)
        set_cell_shading(cell, "E8EEF4")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(data_rows):
        for j, val in enumerate(row):
            if j < len(table.rows[i + 1].cells):
                table.rows[i + 1].cells[j].text = re.sub(r"\*\*([^*]+)\*\*", r"\1", val)
    doc.add_paragraph()


def convert(md_path: Path, out_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines = []
    table_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                block = "\n".join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(block)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Pt(12)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if is_table_row(line):
            cells = parse_table_row(line)
            if not is_separator_row(cells):
                table_buffer.append(cells)
            i += 1
            continue
        elif table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("> "):
            add_formatted_paragraph(doc, stripped[2:], style="Intense Quote")
        elif stripped.startswith("- "):
            add_formatted_paragraph(doc, stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            content = re.sub(r"^\d+\.\s", "", stripped)
            add_formatted_paragraph(doc, content, style="List Number")
        elif stripped == "":
            pass
        else:
            add_formatted_paragraph(doc, stripped)

        i += 1

    if table_buffer:
        add_table(doc, table_buffer)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"wrote {out_path}")


def main():
    targets = DOCS
    if len(sys.argv) > 1:
        md = Path(sys.argv[1])
        out = Path(sys.argv[2]) if len(sys.argv) > 2 else OUT_DIR / (md.stem + ".docx")
        targets = [(md, out)]
    for md_path, out_path in targets:
        convert(md_path, out_path)


if __name__ == "__main__":
    main()
